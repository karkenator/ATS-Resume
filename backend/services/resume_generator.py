"""
Resume generator using python-docx for DOCX and ReportLab for PDF.
Style reference: .claude/resume-style.md

Font: Calibri (DOCX) / Helvetica (PDF).
Default paragraph: Before 0pt, After 4pt, Line spacing 1.08.
Layout: 2 pages. Page 1 = header + skills + first experience. Page 2 = rest + education.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak,
)

from models import GenerateResumeResponse, BulletPoint

# DOCX font
DOCX_FONT = "Calibri"
# PDF fonts (built-in to ReportLab)
PDF_FONT = "Helvetica"
PDF_FONT_BOLD = "Helvetica-Bold"

# Default paragraph spacing
DEFAULT_BEFORE = 0
DEFAULT_AFTER = 4
DEFAULT_LINE = 1.08

# Color constants
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
MEDIUM_GRAY = RGBColor(0x5A, 0x5A, 0x5A)
PERIOD_GRAY = RGBColor(0x59, 0x59, 0x59)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _sp(paragraph, before=DEFAULT_BEFORE, after=DEFAULT_AFTER, line=DEFAULT_LINE):
    """Set paragraph spacing. Defaults to: Before 0pt, After 4pt, Line 1.08."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def _add_bottom_border(paragraph, color="1F4E79"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


def _font(run, name, size, bold=False, color=None, caps=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if caps:
        run.font.all_caps = True


def _add_name(doc, name):
    p = doc.add_paragraph()
    _sp(p, after=0)
    run = p.add_run(name)
    _font(run, DOCX_FONT, 28, color=DARK_BLUE)


def _add_title(doc, title):
    p = doc.add_paragraph()
    _sp(p)
    run = p.add_run(title)
    _font(run, DOCX_FONT, 14, color=MEDIUM_GRAY)


def _add_contact(doc, personal_info):
    parts = [p for p in [personal_info.email, personal_info.phone, personal_info.location] if p]
    urls = [u for u in [personal_info.linkedin_url, personal_info.github_url, personal_info.portfolio_url] if u]

    p = doc.add_paragraph()
    _sp(p, after=0)
    run = p.add_run("  |  ".join(parts))
    _font(run, DOCX_FONT, 10)

    if urls:
        p2 = doc.add_paragraph()
        _sp(p2)
        run2 = p2.add_run("  |  ".join(urls))
        _font(run2, DOCX_FONT, 10)


def _add_section_header(doc, title):
    p = doc.add_paragraph()
    _sp(p, before=6, after=4)
    run = p.add_run(title.upper())
    _font(run, DOCX_FONT, 16, color=DARK_BLUE, bold=True, caps=True)
    _add_bottom_border(p)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
    return p


def _add_summary(doc, summary_text):
    _add_section_header(doc, "Profile")
    p = doc.add_paragraph()
    _sp(p)
    _add_text_with_bold(p, summary_text, font_size=10.5)


def _add_text_with_bold(paragraph, text, font_name=None, font_size=10.5):
    if font_name is None:
        font_name = DOCX_FONT
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _font(run, font_name, font_size, bold=True)
        else:
            run = paragraph.add_run(part)
            _font(run, font_name, font_size)


def _add_skills(doc, skills):
    _add_section_header(doc, "Skills")
    for category, skill_list in skills.items():
        p = doc.add_paragraph()
        _sp(p, after=2)
        run_cat = p.add_run(f"{category}: ")
        _font(run_cat, DOCX_FONT, 10.5, bold=True)
        run_skills = p.add_run(", ".join(skill_list))
        _font(run_skills, DOCX_FONT, 10.5)


def _add_single_experience(doc, exp):
    p_company = doc.add_paragraph()
    _sp(p_company, before=4, after=0)
    run = p_company.add_run(f"{exp.role}, {exp.company_name}")
    _font(run, DOCX_FONT, 12, bold=True)

    p_period = doc.add_paragraph()
    _sp(p_period, after=2)
    run = p_period.add_run(f"{exp.period}  |  {exp.location}")
    _font(run, DOCX_FONT, 11, color=PERIOD_GRAY)

    for bullet in exp.bullets:
        _add_bullet_point(doc, bullet)


def _add_bullet_point(doc, bullet: BulletPoint):
    p = doc.add_paragraph()
    _sp(p, after=2)

    pPr = p._p.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="180"/>')
    pPr.append(ind)

    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="0"/>'
        f'  <w:numId w:val="1"/>'
        f"</w:numPr>"
    )
    pPr.append(numPr)

    if bullet.bold_keywords:
        text = bullet.text
        keywords_sorted = sorted(bullet.bold_keywords, key=len, reverse=True)
        pattern = "|".join(re.escape(kw) for kw in keywords_sorted)
        parts = re.split(f"({pattern})", text, flags=re.IGNORECASE)

        for part in parts:
            if any(part.lower() == kw.lower() for kw in bullet.bold_keywords):
                run = p.add_run(part)
                _font(run, DOCX_FONT, 10.5, bold=True)
            else:
                run = p.add_run(part)
                _font(run, DOCX_FONT, 10.5)
    else:
        run = p.add_run(bullet.text)
        _font(run, DOCX_FONT, 10.5)


def _add_education(doc, education_list):
    _add_section_header(doc, "Education")
    for edu in education_list:
        p_edu = doc.add_paragraph()
        _sp(p_edu, before=4, after=0)
        run = p_edu.add_run(f"{edu.degree}, {edu.institution}")
        _font(run, DOCX_FONT, 12, bold=True)

        p_period = doc.add_paragraph()
        _sp(p_period)
        run = p_period.add_run(f"{edu.start_date} – {edu.end_date}  |  {edu.location}")
        _font(run, DOCX_FONT, 11, color=PERIOD_GRAY)

        if edu.description:
            p_desc = doc.add_paragraph()
            _sp(p_desc)
            run = p_desc.add_run(edu.description)
            _font(run, DOCX_FONT, 10.5)


def _setup_numbering(doc):
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    abstract_num = parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="0">'
        f'  <w:lvl w:ilvl="0">'
        f'    <w:start w:val="1"/>'
        f'    <w:numFmt w:val="bullet"/>'
        f'    <w:lvlText w:val="\u2022"/>'
        f'    <w:lvlJc w:val="left"/>'
        f"    <w:pPr>"
        f'      <w:ind w:left="360" w:hanging="180"/>'
        f"    </w:pPr>"
        f"    <w:rPr>"
        f'      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:hint="default"/>'
        f"    </w:rPr>"
        f"  </w:lvl>"
        f"</w:abstractNum>"
    )
    numbering_xml.append(abstract_num)

    num = parse_xml(
        f'<w:num {nsdecls("w")} w:numId="1">'
        f'  <w:abstractNumId w:val="0"/>'
        f"</w:num>"
    )
    numbering_xml.append(num)


# ---------------------------------------------------------------------------
# DOCX generation — 2-page layout
# ---------------------------------------------------------------------------

def generate_docx(resume_data: GenerateResumeResponse, output_path: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    _setup_numbering(doc)

    # === PAGE 1: Header + Skills + First Experience ===
    pi = resume_data.personal_info
    _add_name(doc, pi.full_name)
    _add_title(doc, pi.title)
    _add_contact(doc, pi)
    _add_summary(doc, resume_data.summary)
    _add_skills(doc, resume_data.skills)

    _add_section_header(doc, "Work Experience")
    if resume_data.experiences:
        _add_single_experience(doc, resume_data.experiences[0])

    # === PAGE 2: Remaining experiences + Education ===
    if len(resume_data.experiences) > 1:
        doc.add_page_break()
        for exp in resume_data.experiences[1:]:
            _add_single_experience(doc, exp)

    _add_education(doc, resume_data.education)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

# PDF leading = fontSize * 1.08 (matching Word's 1.08 line spacing)
def _leading(size):
    return round(size * 1.08 + 4, 1)  # +4pt ≈ After 4pt equivalent


def _bold_keywords_html(text, keywords):
    if not keywords:
        return _escape_html(text)
    keywords_sorted = sorted(keywords, key=len, reverse=True)
    pattern = "|".join(re.escape(kw) for kw in keywords_sorted)
    parts = re.split(f"({pattern})", text, flags=re.IGNORECASE)
    kw_lower = {kw.lower() for kw in keywords}
    html_parts = []
    for part in parts:
        if not part:
            continue
        escaped = _escape_html(part)
        if part.lower() in kw_lower:
            html_parts.append(f"<b>{escaped}</b>")
        else:
            html_parts.append(escaped)
    return "".join(html_parts)


def _markdown_bold_to_html(text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    html_parts = []
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            html_parts.append(f"<b>{_escape_html(part[2:-2])}</b>")
        else:
            html_parts.append(_escape_html(part))
    return "".join(html_parts)


def _escape_html(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ReportLab color constants
RL_DARK_BLUE = HexColor("#1F4E79")
RL_MEDIUM_GRAY = HexColor("#5A5A5A")
RL_PERIOD_GRAY = HexColor("#595959")


def _build_styles():
    s = {}
    s["Name"] = ParagraphStyle(
        "Name", fontName=PDF_FONT, fontSize=28, leading=_leading(28),
        textColor=RL_DARK_BLUE, spaceBefore=0, spaceAfter=0,
    )
    s["Title"] = ParagraphStyle(
        "Title", fontName=PDF_FONT, fontSize=14, leading=_leading(14),
        textColor=RL_MEDIUM_GRAY, spaceBefore=0, spaceAfter=4,
    )
    s["Contact"] = ParagraphStyle(
        "Contact", fontName=PDF_FONT, fontSize=10, leading=_leading(10),
        textColor=black, spaceBefore=0, spaceAfter=0,
    )
    s["ContactLast"] = ParagraphStyle(
        "ContactLast", fontName=PDF_FONT, fontSize=10, leading=_leading(10),
        textColor=black, spaceBefore=0, spaceAfter=4,
    )
    s["SectionHeader"] = ParagraphStyle(
        "SectionHeader", fontName=PDF_FONT_BOLD, fontSize=16, leading=_leading(16),
        textColor=RL_DARK_BLUE, spaceBefore=6, spaceAfter=4,
    )
    s["CompanyRole"] = ParagraphStyle(
        "CompanyRole", fontName=PDF_FONT_BOLD, fontSize=12, leading=_leading(12),
        textColor=black, spaceBefore=4, spaceAfter=0,
    )
    s["Period"] = ParagraphStyle(
        "Period", fontName=PDF_FONT, fontSize=11, leading=_leading(11),
        textColor=RL_PERIOD_GRAY, spaceBefore=0, spaceAfter=2,
    )
    s["SkillLine"] = ParagraphStyle(
        "SkillLine", fontName=PDF_FONT, fontSize=10.5, leading=_leading(10.5),
        textColor=black, spaceBefore=0, spaceAfter=2,
    )
    s["Body"] = ParagraphStyle(
        "Body", fontName=PDF_FONT, fontSize=10.5, leading=_leading(10.5),
        textColor=black, spaceBefore=0, spaceAfter=4,
    )
    s["Bullet"] = ParagraphStyle(
        "Bullet", fontName=PDF_FONT, fontSize=10.5, leading=_leading(10.5),
        textColor=black, leftIndent=14, firstLineIndent=0,
        spaceBefore=0, spaceAfter=2, bulletIndent=0,
    )
    return s


# ---------------------------------------------------------------------------
# PDF generation — 2-page layout
# ---------------------------------------------------------------------------

def generate_pdf(resume_data: GenerateResumeResponse, output_path: Path):
    styles = _build_styles()

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
    )

    story = []
    pi = resume_data.personal_info

    # Header
    story.append(Paragraph(_escape_html(pi.full_name), styles["Name"]))
    story.append(Paragraph(_escape_html(pi.title), styles["Title"]))

    contact_parts = [p for p in [pi.email, pi.phone, pi.location] if p]
    story.append(Paragraph(_escape_html("  |  ".join(contact_parts)), styles["Contact"]))
    url_parts = [p for p in [pi.linkedin_url, pi.github_url, pi.portfolio_url] if p]
    if url_parts:
        story.append(Paragraph(_escape_html("  |  ".join(url_parts)), styles["ContactLast"]))

    # === PAGE 1: Profile + Skills + First Experience ===
    _pdf_section_header(story, "PROFILE", styles)
    summary_html = _markdown_bold_to_html(resume_data.summary)
    story.append(Paragraph(summary_html, styles["Body"]))

    _pdf_section_header(story, "SKILLS", styles)
    for category, skill_list in resume_data.skills.items():
        cat_html = f"<b>{_escape_html(category)}:</b> {_escape_html(', '.join(skill_list))}"
        story.append(Paragraph(cat_html, styles["SkillLine"]))

    _pdf_section_header(story, "WORK EXPERIENCE", styles)
    if resume_data.experiences:
        exp = resume_data.experiences[0]
        story.append(Paragraph(
            _escape_html(f"{exp.role}, {exp.company_name}"), styles["CompanyRole"]
        ))
        story.append(Paragraph(
            _escape_html(f"{exp.period}  |  {exp.location}"), styles["Period"]
        ))
        for bullet in exp.bullets:
            bullet_html = _bold_keywords_html(bullet.text, bullet.bold_keywords)
            story.append(Paragraph(
                f"<bullet>\u2022</bullet> {bullet_html}", styles["Bullet"]
            ))

    # === PAGE 2: Remaining experiences + Education ===
    if len(resume_data.experiences) > 1:
        story.append(PageBreak())
        for exp in resume_data.experiences[1:]:
            story.append(Paragraph(
                _escape_html(f"{exp.role}, {exp.company_name}"), styles["CompanyRole"]
            ))
            story.append(Paragraph(
                _escape_html(f"{exp.period}  |  {exp.location}"), styles["Period"]
            ))
            for bullet in exp.bullets:
                bullet_html = _bold_keywords_html(bullet.text, bullet.bold_keywords)
                story.append(Paragraph(
                    f"<bullet>\u2022</bullet> {bullet_html}", styles["Bullet"]
                ))

    _pdf_section_header(story, "EDUCATION", styles)
    for edu in resume_data.education:
        story.append(Paragraph(
            _escape_html(f"{edu.degree}, {edu.institution}"), styles["CompanyRole"]
        ))
        story.append(Paragraph(
            _escape_html(f"{edu.start_date} \u2013 {edu.end_date}  |  {edu.location}"),
            styles["Period"],
        ))
        if edu.description:
            story.append(Paragraph(_escape_html(edu.description), styles["Body"]))

    doc.build(story)


def _pdf_section_header(story, title, styles):
    story.append(Paragraph(title, styles["SectionHeader"]))
    story.append(HRFlowable(
        width="100%", thickness=0.5, color=RL_DARK_BLUE,
        spaceBefore=0, spaceAfter=2,
    ))
